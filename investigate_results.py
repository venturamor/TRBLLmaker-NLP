import pandas as pd
import os
import sys
import datasets
import nltk
from nltk.tokenize import word_tokenize
# from nltk.corpus import stopwords
import docx
import numpy as np
import datetime
from box import Box
import yaml
from tqdm import tqdm
from prompts import *
from config_parser import *
import matplotlib.pyplot as plt


def generate_graphs(df, pickl_path, std_df=None):
    """
    Generate graphs from the dataframe
    """
    df_unstacked = df.unstack(level=0)
    # with err bars - yerr - need to take from std
    if std_df!=None:
        ax = df_unstacked.plot(kind='bar', rot=0, figsize=(9, 7), layout=(2, 3), yerr=0.25)

    ax = df_unstacked.plot(kind='bar', rot=0, figsize=(9, 7), layout=(2, 3))
    # set x-axis as df.index.names[1]
    if len(df.index.names) > 1:
        ax.set_xlabel(df.index.names[1])
        # set title
        ax.set_title("{} score for different {} and {}".format(df.keys()[0], df.index.names[0], df.index.names[1]))
    else:
        ax.set_xlabel(df.index.names[0])
        # set title
        ax.set_title("{} score for different {}".format(df.keys()[0], df.index.names[0]))
    # set y-axis as df.keys()[0]
    ax.set_ylabel(df.keys()[0])
    # change Legends names
    names = [name for name in df.index.get_level_values(0).unique()]
    ax.legend(loc='upper center', labels=names)
    # save
    main_path = private_args.path.main_path
    plot_path = os.path.join(main_path, 'plots', pickl_path.split('/')[-2])
    if not os.path.exists(plot_path):
        os.makedirs(plot_path)
    if len(df.index.names) > 1:
        plt.savefig(os.path.join(plot_path, '{}_{}_{}.png'.format(df.keys()[0], df.index.names[0], df.index.names[1])))
    else:
        plt.savefig(os.path.join(plot_path, '{}_{}.png'.format(df.keys()[0], df.index.names[0])))
    plt.show()
    # merge columns with the same name
    # def get_notnull(x): return ';'.join(x[x.notnull()].astype(str))

    # df_clean = df.groupby(level=0, axis=1).apply(lambda x: x.apply(get_notnull, axis=1))

    # change dataframe values from string to int
    # df_clean = df_clean.apply(pd.to_numeric, errors='coerce')

    # models = ['gpt2', 'gpt2-medium', 'EleutherAI/gpt-neo-1.3B']
    # prompts = ['lyrics_meaning', 'lyrics_meaning_with_metadata', 'song', 'song_with_metadata',
    # 'question_context', 'question_context_with_metadata', None]

    # plot bar plots

def combine_before_and_after(pickle_before, pickle_after, after_name):
    df_before = pd.read_pickle(pickle_before)
    df_after = pd.read_pickle(pickle_after)
    # take only relevant columns
    df_before_narrowed = df_before[['decode_method','example_index', 'input_prompt', 'prompt_type', 'model', 'gt_meaning', 'predicted_meaning']]
    after_prompt = df_after['prompt_type'][0]
    df_before_narrowed = df_before_narrowed[df_before_narrowed['prompt_type'] == after_prompt]

    df_before_narrowed = df_before_narrowed[df_before_narrowed['model'] == 'gpt2-medium']

    df_after_narrowed = df_after[['predicted_meaning', 'example_index']]
    # for each example_index in df_before_narrowed, find the corresponding example_index in df_after_narrowed
    # and add the predicted_meaning to the df_before_narrowed
    for index, row in df_before_narrowed.iterrows():
        example_index = row['example_index']
        prediction_after = df_after_narrowed.loc[df_after_narrowed['example_index'] == example_index,
                              'predicted_meaning']
        df_before_narrowed.loc[index, 'predicted_meaning_after'] = prediction_after.values[0]

    # print result to docx file
    doc = docx.Document()
    doc.add_heading('Before and after', 0)
    doc.add_paragraph('The following table shows the predicted meaning before and after training.')
    for index, row in df_before_narrowed.iterrows():
        paragraph = \
            doc.add_paragraph("input prompt:\n{}\n\ngt meaning:\n{}\n\ndecode method:\n{}\n\n".format(
                                    row['input_prompt'],
                                    row['gt_meaning'],
                                    row['decode_method']))
        # highlight the predicted meaning
        paragraph.add_run("predicted meaning before training:\n{})".format(
            row['predicted_meaning'])).bold = True
        paragraph.add_run("\n\npredicted meaning after training:\n{})".format(
            row['predicted_meaning_after'])).font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW

    after_folder = os.path.join(private_args.path.main_path, 'final_results')

    if not os.path.exists(after_folder):
        os.makedirs(after_folder)
    doc.save(os.path.join(after_folder, 'before_and_after{}.docx'.format(after_name)))



def get_scores_after_training():
    pickle_list = ["analysis_rouge1_['decode_method', 'model'].pkl", "analysis_total_score_['decode_method', 'model'].pkl",
                   "analysis_cos_pred_label_['decode_method', 'model'].pkl"]
    pickles_path = '/home/tok/TRBLLmaker/post_eval/analysis_results'
    # iterate over all folders in pickles_path
    for folder in os.listdir(pickles_path):
        if 'lyrics_meaning_with_metadata' in folder \
            or 'question_context_with_metadata' in folder:
            # iterate over all pickles in pickle_list
            for pickle_path in pickle_list:
                pickle_path = os.path.join(pickles_path, folder, pickle_path)
                df = pd.read_pickle(pickle_path)
                print(1)


                # df = pd.read_pickle(os.path.join(pickles_path, folder, pickle_path))
                # df_unstacked = df.unstack(level=0)
                # # plot bar plot
                # ax = df_unstacked.plot(kind='bar', rot=0, figsize=(9, 7), layout=(2, 3))
                # if len(df.index.names) > 1:
                #     ax.set_xlabel(df.index.names[1])
                #     # set title
                #     ax.set_title(
                #         "{} score for different {} and {}".format(df.keys()[0], df.index.names[0], df.index.names[1]))
                # else:
                #     ax.set_xlabel(df.index.names[0])
                #     # set title
                #     ax.set_title("{} score for different {}".format(df.keys()[0], df.index.names[0]))
                # # set y-axis as df.keys()[0]
                # ax.set_ylabel(df.keys()[0])
                # # change Legends names
                # names = [name for name in df.index.get_level_values(0).unique()]
                # ax.legend(loc='upper center', labels=names)
                # plt.show()

if __name__ == '__main__':
    # get_scores_after_training()
    # exit()
    after_training_folder = '/home/tok/TRBLLmaker/results/after_training/predictions_after_training'
    before_training_folder = '/home/tok/TRBLLmaker/results/pretraining/predictions_before_training/before_training'
    pickle_name = 'inference_results.pkl'
    pickle_before = os.path.join(before_training_folder, pickle_name)
    # iterate over all folders in after_training_folder
    for folder in os.listdir(after_training_folder):
        # new docx file
        # doc = docx.Document()
        pickle_after = os.path.join(after_training_folder, folder, pickle_name)
        after_name = folder.split('trained_model_checkpoint_')[-1]
    # read pickle
        df_before = pd.read_pickle(pickle_before)
        df_after = pd.read_pickle(pickle_after)
        # sort by example_index
        df_before = df_before.sort_values(by='example_index')
        df_after = df_after.sort_values(by='example_index')
        print(1)
        # # print df_after to docx
        # paragraph = doc.add_paragraph("example_num: \n{}\n\n".format(df_after['example_index'].values))
        # paragraph.add_run("input: \n{}\n\n".format(df_after['input_prompt'].values))
        # paragraph.add_run("decode_method: \n{}\n\n".format(df_after['decode_method'].values))
        # # next line in bold
        # paragraph.add_run("predicted_meaning: \n{}\n\n".format(df_after['predicted_meaning'].values)).bold = True
        # # save docx
        # doc.save(os.path.join(after_training_folder, folder, 'before_and_after{}_{}.docx'.format(after_name, pickle_name)))
        # print into docx
        # combine_before_and_after(pickle_before, pickle_after, after_name=after_name)
    # exit()

    # main_path = private_args.path.main_path
    # eval_path = "post_eval"
    # after_folder = '/home/tok/TRBLLmaker/post_eval/post_eval/analysis_results' #training_args.path_args.pretraining_folder #'before_training'
    # before_folder = '/home/tok/TRBLLmaker/post_eval/post_eval/pre_training' #training_args.path_args.after_training_folder #'after_training'
    # results_folder = training_args.path_args.results_path
    # Load pickle as a dataframe
    # pickle_name = "predictions_before_training_2022-03-12-17-41-00.pkl"
    # # if pickel name has 'before' in it, load before pickle
    # if 'before' in pickle_name:
    #     folder = before_folder
    # else:
    #     folder = after_folder
    # pickles_folder = os.path.join(private_args.path.main_path, results_folder, folder)
    #
    # pickles_path = '/home/tok/TRBLLmaker/post_eval/post_eval/pre_training'

    # # combine before and after pickles
    # pickle_before = r'/home/tok/TRBLLmaker/results/pretraining/predictions_before_training_2022-03-12-17-41-00.pkl'
    # pickle_after =r'/home/tok/TRBLLmaker/results/after_training/predictions_after_training_2022-03-14-11-17-57.pkl'
    # combine_before_and_after(pickle_before, pickle_after)

    # pickl_name = "analysis_cos_pred_label_['decode_method', 'model', 'prompt_type']_2022-03-12 19:37:39.523430.pkl"#"full_analysis_120322_1939.pkl" # full_analysis_120322_1937.pkl
    # pickls_path = os.path.join(main_path, eval_path)
    # folder_path = '/home/tok/TRBLLmaker/post_eval/post_eval/analysis_results'
    # # iterate over al folders if folder_path
    # for pickles_path in os.listdir(folder_path):
    #     # iterate over all pickles in pickle_path
    #     for pickle_name in os.listdir(pickles_path):
    #         if pickle_name.endswith(".pkl") and pickle_name.startswith("analysis")\
    #                 and 'std' not in pickle_name and pickle_name != 'analysis_total_score.pkl':
    #             # load pickle
    #             pickl_path = os.path.join(pickles_path, pickle_name)
    #             df = pd.read_pickle(pickl_path)
    #             # generate graphs
    #             generate_graphs(df, pickl_path)

    # compare between models
    # bar plot for all models and compare all score.
    # x-axis represents the prompts, y-axis represents the scores values
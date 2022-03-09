import os

from transformers import GPT2LMHeadModel, GPT2Tokenizer,  GPTNeoForCausalLM
import datasets
import torch
import docx
import datetime
from transformers import T5Tokenizer, T5ForConditionalGeneration, T5Model, T5TokenizerFast, GPT2Tokenizer, \
    GPT2LMHeadModel, Trainer, GPTNeoForCausalLM, TrainingArguments, TFGPT2LMHeadModel
from transformers import DataCollatorForSeq2Seq
from box import Box
import yaml
from config_parser import *
from tqdm import tqdm
from prompts import *
import pandas as pd
import numpy as np


def run_inference_on_sample(model_name, input_prompt, decode_method_index=1, temperature=0.9, num_return_sequence=1,
                            TF=False):
    tokenizer = GPT2Tokenizer.from_pretrained(model_name)
    if model_name == 'gpt2' or model_name =='gpt2-medium':
        if TF:
            model = TFGPT2LMHeadModel.from_pretrained(model_name)
        else:
            model = GPT2LMHeadModel.from_pretrained(model_name)
    else:
        model = GPTNeoForCausalLM.from_pretrained(model_name)
    # print("Model: {} loaded".format(model_name))

    df_inference = pd.DataFrame(columns=['input_prompt', 'predicted_text', 'decode_method', 'temperature'])

    decode_methods = ['greedy', 'beam search', 'sampling', 'top-k sampling', 'top-p sampling']
    for decode_method in decode_methods:
        # encode prompt
        if TF:
            input_ids = tokenizer.encode(input_prompt, return_tensors='tf')
        else:
            input_ids = tokenizer(input_prompt, return_tensors="pt").input_ids

        # predict
        if decode_method == 'greedy':
            # greedy
            outputs = model.generate(input_ids, max_length=training_args.eval_pretrained_args.max_length,
                                     temperature=temperature, num_return_sequence=num_return_sequence,
                                     num_return_sequences=1,
                                     )
        elif decode_method == 'beam search':
            # beam search with penalty on repeat
            outputs = model.generate(input_ids, max_length=training_args.eval_pretrained_args.max_length,
                                     num_beams=3, early_stopping=True,
                                     num_return_sequences=training_args.eval_pretrained_args.num_return_sequences,
                                     no_repeat_ngram_size=2
                                     )
        elif decode_method == 'sampling':
            # sampling
            outputs = model.generate(input_ids, do_sample=True, top_k=0,
                                     max_length=training_args.eval_pretrained_args.max_length,
                                     num_return_sequences=training_args.eval_pretrained_args.num_return_sequences,
                                     temperature=temperature
                                     )
        elif decode_method == 'top-k sampling':
            # top-k sampling
            outputs = model.generate(input_ids, do_sample=True, top_k=50,
                                     max_length=training_args.eval_pretrained_args.max_length,
                                     num_return_sequences=training_args.eval_pretrained_args.num_return_sequences,
                                     )
        else:  # decode_method == 'top-p sampling' (default)
            # top-p sampling
            outputs = model.generate(input_ids, do_sample=True, top_k=0, top_p=0.92,
                                     max_length=training_args.eval_pretrained_args.max_length,
                                     num_return_sequences=training_args.eval_pretrained_args.num_return_sequences,
                                     )
        # additional parameters for better decoding - repetition_penalty, min_length
        # decode
        pred_text = tokenizer.batch_decode(outputs, skip_special_tokens=True)

        input_list, generated_text = [], []

        for pred in pred_text:
            input_list.append(input_prompt)
            generated_text.append(pred)

        df_curr = pd.DataFrame({'input_prompt': input_list, 'predicted_text': generated_text,
                                 'decode_method': decode_method, 'temperature': temperature})
        df_inference = pd.concat([df_inference, df_curr])
    return df_inference


def compare_models(models_names, file_name, TF=False):
    """
    A function to compare the performance of different models.
    """
    prompt_types = ['lyrics_meaning', 'lyrics_meaning_with_metadata', 'song', 'song_with_metadata',
                    'question_context', 'question_context_with_metadata', None]

    # Variables
    max_input_length = 512
    max_target_length = 512
    temperature = 0.9
    N = 2
    num_return_sequences = 2
    full_df = pd.DataFrame(columns=['input_prompt', 'predicted_text', 'decode_method', 'temperature',
                                    'model', 'prompt_type'])
    # create a doc file to write the generated prompts
    doc = docx.Document()
    doc.add_heading('Predicted annotations by different models, prompts and temperature', 0)

    # load datasets
    dataset_name = 'TRBLL_dataset.py'
    samples_dataset = datasets.load_dataset(dataset_name)['test']
    print("Loaded {} samples from {}".format(len(samples_dataset), dataset_name))

    # set the seed for reproducibility
    torch.manual_seed(21)
    np.random.seed(21)

    #  Run for each model
    for model_name in tqdm(models_names):
        print("Run for each model.\nCurrent model: {}".format(model_name))
        # load the model
        tokenizer = GPT2Tokenizer.from_pretrained(model_name)
        if model_name == 'gpt2' or model_name == 'gpt2-medium':
            if TF:
                model = TFGPT2LMHeadModel.from_pretrained(model_name)
            else:
                model = GPT2LMHeadModel.from_pretrained(model_name)
        else:
            model = GPTNeoForCausalLM.from_pretrained(model_name)

        # choose random N samples from the dataset
        samples = torch.randint(0, len(samples_dataset['data']) - 1, (N,))

        # Run for each sample
        for index in tqdm(samples):
            print("Run for each sample...")
            lyrics = samples_dataset['data'][index][0]
            meaning = samples_dataset['labels'][index][0]
            artist = samples_dataset['artist'][index][0]
            title = samples_dataset['title'][index][0]

            # Run for each prompt type
            for prompt_type in tqdm(prompt_types):
                print("Run for each prompt type.\nCurrent prompt type:{}".format(prompt_type))
                input_prompt = generate_prompts(lyrics=lyrics, meaning=meaning, artist=artist, title=title,
                                                prompt_type=prompt_type)

                temperature = 0.9
                num_return_sequence = 1,
                TF = False
                df_inference = pd.DataFrame(columns=['input_prompt', 'predicted_text', 'decode_method', 'temperature'])

                decode_methods = ['greedy', 'beam search', 'sampling', 'top-k sampling', 'top-p sampling']
                for decode_method in decode_methods:
                    # encode prompt
                    if TF:
                        input_ids = tokenizer.encode(input_prompt, return_tensors='tf')
                    else:
                        input_ids = tokenizer(input_prompt, return_tensors="pt").input_ids

                    # predict
                    if decode_method == 'greedy':
                        # greedy
                        outputs = model.generate(input_ids, max_length=training_args.eval_pretrained_args.max_length,
                                                 temperature=temperature, num_return_sequence=num_return_sequence,
                                                 num_return_sequences=1,
                                                 )
                    elif decode_method == 'beam search':
                        # beam search with penalty on repeat
                        outputs = model.generate(input_ids, max_length=training_args.eval_pretrained_args.max_length,
                                                 num_beams=3, early_stopping=True,
                                                 num_return_sequences=training_args.eval_pretrained_args.num_return_sequences,
                                                 no_repeat_ngram_size=2
                                                 )
                    elif decode_method == 'sampling':
                        # sampling
                        outputs = model.generate(input_ids, do_sample=True, top_k=0,
                                                 max_length=training_args.eval_pretrained_args.max_length,
                                                 num_return_sequences=training_args.eval_pretrained_args.num_return_sequences,
                                                 temperature=temperature
                                                 )
                    elif decode_method == 'top-k sampling':
                        # top-k sampling
                        outputs = model.generate(input_ids, do_sample=True, top_k=50,
                                                 max_length=training_args.eval_pretrained_args.max_length,
                                                 num_return_sequences=training_args.eval_pretrained_args.num_return_sequences,
                                                 )
                    else:  # decode_method == 'top-p sampling' (default)
                        # top-p sampling
                        outputs = model.generate(input_ids, do_sample=True, top_k=0, top_p=0.92,
                                                 max_length=training_args.eval_pretrained_args.max_length,
                                                 num_return_sequences=training_args.eval_pretrained_args.num_return_sequences,
                                                 )
                    # additional parameters for better decoding - repetition_penalty, min_length
                    # decode
                    pred_text = tokenizer.batch_decode(outputs, skip_special_tokens=True)

                    input_list, generated_text = [], []

                    for pred in pred_text:
                        input_list.append(input_prompt)
                        generated_text.append(pred)

                    df_curr = pd.DataFrame({'input_prompt': input_list, 'predicted_text': generated_text,
                                            'decode_method': decode_method, 'temperature': temperature,
                                            'model': model_name, 'prompt_type': prompt_type})
                    df_inference = pd.concat([df_inference, df_curr], ignore_index=True)

                full_df = pd.concat([full_df, df_inference], ignore_index=True)
                for i, row in full_df.iterrows():
                    generated, input_text = row['predicted_text'], row['input_prompt']
                    # Save to docx file
                    para = doc.add_paragraph("Model: {},\n prompt: {},\n temperature: {} \n\n"
                                             .format(model_name, prompt_type, temperature))
                    para.add_run("lyrics: {}.\n meaning: {} \n\n".format(lyrics, meaning))
                    # Print the generated prompt highlighted with green color
                    para.add_run("Gerenated text:\n").font.highlight_color \
                        = docx.enum.text.WD_COLOR_INDEX.RED
                    para.add_run("{} ".format(input_prompt)).font.highlight_color = \
                        docx.enum.text.WD_COLOR_INDEX.YELLOW
                    para.add_run("{} \n\n\n".format("EMPTY" if len(generated.split(input_prompt)) <= 1
                                                    else generated.split(input_prompt)[1])).font.highlight_color = \
                        docx.enum.text.WD_COLOR_INDEX.GREEN
    doc.save('{}_{}.docx'.format(file_name, datetime.datetime.now().strftime("%Y-%m-%d-%H-%M-%S")))
    # save df as pickle
    full_df.to_pickle('{}_{}.pkl'.format(file_name, datetime.datetime.now().strftime("%Y-%m-%d-%H-%M-%S")))
    print("Saved {}_{}.pkl".format(file_name, datetime.datetime.now().strftime("%Y-%m-%d-%H-%M-%S")))
    # Save to csv file
    full_df.to_csv('{}_{}.csv'.format(file_name, datetime.datetime.now().strftime("%Y-%m-%d-%H-%M-%S")))
    print("Saved {}_{}.csv".format(file_name, datetime.datetime.now().strftime("%Y-%m-%d-%H-%M-%S")))
    print("Done")


if __name__ == '__main__':
    models_names = ['gpt2', 'gpt2-medium', 'EleutherAI/gpt-neo-1.3B', 'EleutherAI/gpt-neo-2.7B']
    main_path = private_args.path.main_path
    results_path = private_args.path.results_path
    pretraining_folder = private_args.path.pretraining_folder
    file_name = "predictions_before_training"
    file_path = os.path.join(main_path, results_path, pretraining_folder, file_name)
    compare_models(models_names, file_name=file_path)

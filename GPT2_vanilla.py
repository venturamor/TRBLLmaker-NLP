from transformers import GPT2LMHeadModel, GPT2Tokenizer,  GPTNeoForCausalLM
import datasets
import torch
import docx
import datetime
from transformers import T5Tokenizer, T5ForConditionalGeneration, T5Model, T5TokenizerFast, GPT2Tokenizer, \
    GPT2LMHeadModel, Trainer, GPTNeoForCausalLM, TrainingArguments
from transformers import DataCollatorForSeq2Seq

def generate_prompts(lyrics, meaning, artist, title, prompt_type):
    if prompt_type == "lyrics_meaning":
        data = "lyrics: {}.\n meaning:".format(lyrics)
    elif prompt_type == "song_metadata":
        # Load the songs and annotations
        data = 'Explain the song "{}", written by {}.\n Lyrics: {}.\n Explanation:'.format(title, artist, lyrics)
    elif prompt_type == "question_context":
        data = 'question: what is the meaning of {} in the song "{}", ' \
               'written by {}?\n context: {}.\n answer:'.format(meaning, title, artist, lyrics)
    else:  # None: no prompt
        data = lyrics
    return data


models_names = ['gpt2']#'EleutherAI/gpt-neo-1.3B', 'EleutherAI/gpt-neo-2.7B']  # 'gpt2-medium'
prompt_types = ['lyrics_meaning', 'song_metadata', 'question_context']
max_input_length = 512
max_target_length = 128
temperature = 0.9
N = 2
num_return_sequences = 2
torch.manual_seed(21)

for model_name in models_names:
    tokenizer = GPT2Tokenizer.from_pretrained(model_name)
    if model_name == 'gpt2' or model_name =='gpt2-medium':
        model = GPT2LMHeadModel.from_pretrained(model_name)
    else:
        model = GPTNeoForCausalLM.from_pretrained(model_name)
    print("Model: {} loaded".format(model_name))

    # load datasets
    dataset_name = 'TRBLL_dataset.py'
    samples_dataset = datasets.load_dataset(dataset_name)['train']
    print("Loaded {} samples from {}".format(len(samples_dataset), dataset_name))

    # choose random N samples from the dataset
    samples = torch.randint(0, len(samples_dataset['data']) - 1, (N,))

    # create a doc file to write the generated prompts
    doc = docx.Document()
    doc.add_heading('Predicted annotations by different models, prompts and temperature', 0)

    for index in samples:
        lyrics = samples_dataset['data'][index][0]
        meaning = samples_dataset['labels'][index][0]
        artist = samples_dataset['artist'][index][0]
        title = samples_dataset['title'][index][0]

        for prompt_type in prompt_types:
            print("Generating prompts for {}".format(prompt_type))
            input_prompt = generate_prompts(lyrics, meaning, artist, title, prompt_type)
            input_ids = tokenizer(input_prompt, return_tensors="pt").input_ids

            print("Generating {} prompts for {}".format(num_return_sequences, input_prompt))

            gen_tokens = model.generate(
                input_ids,
                do_sample=True,
                temperature=temperature,
                max_length=max_target_length,
                num_return_sequences=num_return_sequences,
                # top_k=50,
                # top_p=0.95,
            )
            generated_list = tokenizer.batch_decode(gen_tokens)
            print("Generated prompt: {}".format(input_prompt))
            for i, sample_output in enumerate(gen_tokens):
                gen_text = tokenizer.decode(sample_output, skip_special_tokens=True)
                # Save to docx file
                para = doc.add_paragraph("Model: {}, prompt: {}, temperature: {} \n\n"
                                         .format(model_name, prompt_type, temperature))
                para.add_run("lyrics: {}. meaning: {} \n\n".format(lyrics, meaning))
                # Print the generated prompt highlighted with green color
                for generated in generated_list:
                    para.add_run("Gerenated text:\n").font.highlight_color = docx.enum.text.WD_COLOR_INDEX.RED
                    para.add_run("{} \n\n\n".format(generated)).font.highlight_color = docx.enum.text.WD_COLOR_INDEX.GREEN
                    print("Generated text: {}".format(generated))
    doc.save('predictions_before_training_{}.docx'.format(datetime.datetime.now().strftime("%Y-%m-%d-%H-%M-%S")))
    print("Predictions saved to file")
    print("Done")

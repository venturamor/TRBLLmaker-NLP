import pandas as pd
import os
import sys
import datasets
import nltk
from nltk.tokenize import word_tokenize
# from nltk.corpus import stopwords
import docx
import numpy as np
import datetime
from box import Box
import yaml
from tqdm import tqdm
from prompts import *
from config_parser import *


# nltk.download()

def calc_rouge(sen_a, sen_b):
    rouge = datasets.load_metric('rouge')
    # if one of the sentences is empty, return 0
    if sen_a == "" or sen_b == "":
        rouge_score = rouge.compute(predictions="1", references="0")
    else:
        # fix mismatch by length by cut:
        # match in the number of len prediction and len reference :len(sen_b)
        if len(sen_a) >= len(sen_b):
            rouge_score = rouge.compute(predictions=sen_a[:len(sen_b)], references=sen_b)
        else:
            rouge_score = rouge.compute(predictions=sen_a, references=sen_b[:len(sen_a)])
    # low, mid, high -  """Tuple containing confidence intervals for scores."""
    precentile = 1  # 'mid'
    score_type = 2  # 'fmeasure'  # recall, precision
    rouge1 = rouge_score['rouge1'][precentile][score_type]
    rouge2 = rouge_score['rouge2'][precentile][score_type]
    return rouge1, rouge2


def calc_cosine_similarity_2_sentences(sen_a, sen_b):
    """
    calculates cosine similarity of 2 sentences -
    counting vector (histogram) per each word for each sentence
    :param sen_a:
    :param sen_b:
    :return:
    """
    # splits to words
    a_list = word_tokenize(sen_a)
    b_list = word_tokenize(sen_b)
    # remove stop words
    # sw = stopwords.words('english')
    # a_set = {w for w in a_list if not w in sw}
    # b_set = {w for w in b_list if not w in sw}
    a_set = {w for w in a_list}
    b_set = {w for w in b_list}

    l1 = []
    l2 = []
    # form a set containing keywords of both strings
    rvector = a_set.union(b_set)  # all words
    for w in rvector:
        if w in a_set:
            l1.append(1)  # create a vector
        else:
            l1.append(0)
        if w in b_set:
            l2.append(1)
        else:
            l2.append(0)
    c = 0
    # cosine formula
    for i in range(len(rvector)):
        c += l1[i] * l2[i]

    try:
        cosine = c / float((sum(l1) * sum(l2)) ** 0.5)
    except:
        cosine = 0
    return cosine


# todo: remove splitting the prediction - duplicated ( already done in funetunings_scripts)
def fix_columns(df):
    predicted_meaning = []
    gt_meaning = []
    for in_prompt, pred in zip(df['input_prompt'], df['predicted_text']):
        pred_splitted = pred.split(in_prompt)

        if len(pred_splitted) <= 1:
            pred = "Empty"
        elif len(pred_splitted) == 2:
            pred = pred_splitted[1]
        # else:
        #     pred = "More than one repetition: " + pred
        predicted_meaning.append(pred)

    df['predicted_meaning'] = predicted_meaning

    df['gt_meaning'] = predicted_meaning  # TODO: delete
    df['lyrics'] = df['input_prompt']  # TODO: delete
    return df


def post_eval(pickle_path, fix_flag=0):
    # Load pickle as a dataframe
    df = pd.read_pickle(pickle_path)
    # pickle_name = os.path.split(pickle_path)[1]
    pickle_name = pickle_path.split('/')[-2] + datetime.datetime.today().strftime('_%d%m%y_%H%M')

    if fix_flag:
        df = fix_columns(df)

    # calculate eval_metrices - (input, prediction), (label, prediction)
    cos_pred_lyrics_l, cos_pred_label_l, rouge1_l, rouge2_l = [], [], [], []
    total_score_l = []
    weights_per_metric = {'cos_pred_lyrics': 0.5, 'cos_pred_label': 0.5, 'rouge1': 0.5, 'rouge2': 0}
    for lyrics, pred, label in tqdm(zip(df['lyrics'], df['predicted_meaning'], df['gt_meaning'])):
        # cosine similarity - LSA
        cos_pred_lyrics = calc_cosine_similarity_2_sentences(pred, lyrics)
        cos_pred_label = calc_cosine_similarity_2_sentences(pred, label)
        # rouge
        rouge1, rouge2 = calc_rouge(pred, label)

        scores = {'cos_pred_lyrics': cos_pred_lyrics, 'cos_pred_label': cos_pred_label, 'rouge1': rouge1,
                  'rouge2': rouge2}
        weighted_scores = [scores[k] * v for k, v in weights_per_metric.items()]
        total_score = max(0, sum(weighted_scores) - 2 * weighted_scores[0])  # sum of all minus similarity to lyrics

        # appends
        cos_pred_lyrics_l.append(cos_pred_lyrics)
        cos_pred_label_l.append(cos_pred_label)
        rouge1_l.append(rouge1)
        rouge2_l.append(rouge2)
        total_score_l.append(total_score)

    df['rouge1'] = rouge1_l
    df['rouge2'] = rouge2_l
    df['cos_pred_label'] = cos_pred_label_l
    df['cos_pred_lyrics'] = cos_pred_lyrics_l
    df['total_score'] = total_score_l
    new_pickle_path = "post_eval/example_to_analysis_" + pickle_name + ".pkl"
    df.to_pickle(new_pickle_path)
    print('done')
    return df, new_pickle_path


def analysis(df: pd.DataFrame, compare_params: list, score_name: str,
             new_pickle_path, post_eval_path, run_all=1):
    """

    :param df:
    :param compare_param: list of strings - order - hirrechy - column - model, prompt, decode, any other...
    :return:
    """

    df = pd.read_pickle(new_pickle_path) #TODO: remove
    df_analysis = pd.DataFrame()
    df_analysis_std = pd.DataFrame()

    # creat docx file
    doc = docx.Document()
    doc.add_heading('Analysis', 0)

    if run_all:
        compare_params_lists = [['model', 'prompt_type', 'decode_method'],
                                ['prompt_type', 'model',  'decode_method'],
                                ['decode_method', 'model', 'prompt_type']]
        score_name_list = ['total_score', 'rouge1', 'cos_pred_label', 'cos_pred_lyrics']

        for compare_list in compare_params_lists:
            for score in score_name_list:
                para = doc.add_paragraph('Compare_list: \n{}\nMean score by:{}'.format(compare_list, score))
                print('Compare_list:', compare_list, '\nMean score by:', score)
                h = len(compare_list)  # hierarchies
                for ind_param in range(h):
                    gk = df.groupby(compare_list[:ind_param + 1])
                    mean_gk = gk[score].mean()
                    std_gk = gk[score].std()
                    # save as df
                    mean_gk_df = mean_gk.to_frame()
                    # TODO: add std
                    std_gk_df = std_gk.to_frame()

                    # concat to analysis df
                    df_analysis = pd.concat([df_analysis, mean_gk_df], axis=1)
                    df_analysis_std = pd.concat([df_analysis_std, std_gk_df], axis=1)

                    # save pickle
                    curr_pickle_path = os.path.join(post_eval_path, 'analysis_{}_{}.pkl'.format(
                        score, compare_params[:ind_param + 1]))

                    mean_gk_df.to_pickle(curr_pickle_path)

                    curr_std_pickle_path = os.path.join(post_eval_path, 'analysis_std_{}_{}.pkl'.format(
                        score, compare_params[:ind_param + 1]))

                    std_gk_df.to_pickle(curr_std_pickle_path)

                    # append to docx
                    para.add_run('Mean Hierarchy \n{}:\n'.format(ind_param))
                    para.add_run('Mean:\n{}\n'.format(mean_gk))
                    para.add_run('\n')

                    print('Mean Hierarchy {}:\n'.format(ind_param), mean_gk)

    else:
        h = len(compare_params)  # hierarchies
        for ind_param in range(h):
            gk = df.groupby(compare_params[:ind_param+1])
            mean_gk = gk[score_name].mean()
            std_gk = gk[score_name].std()

            # save as df
            mean_gk_df = mean_gk.to_frame()
            std_gk_df = std_gk.to_frame()

            # concat to analysis df
            df_analysis = pd.concat([df_analysis, mean_gk_df], axis=1)
            df_analysis_std = pd.concat([df_analysis_std, std_gk_df], axis=1)

            # save pickle
            curr_pickle_path = os.path.join(post_eval_path, 'analysis_{}_{}.pkl'.format(
                score_name, compare_params[:ind_param+1]))
            mean_gk_df.to_pickle(curr_pickle_path)

            curr_std_pickle_path = os.path.join(post_eval_path, 'analysis_std_{}_{}.pkl'.format(
                score_name, compare_params[:ind_param + 1]))

            std_gk_df.to_pickle(curr_std_pickle_path)
            # append to docx
            doc.add_paragraph('Mean:\n{}\n'.format(mean_gk))

            print('Mean:\n', mean_gk)

    # save the docx file
    doc.save(os.path.join(post_eval_path, 'analysis_{}.docx'.format(score_name)))

    # save pickle
    df_analysis.to_pickle(os.path.join(post_eval_path, 'analysis_{}.pkl'.format(score_name)))
    df_analysis_std.to_pickle(os.path.join(post_eval_path, 'analysis_std_{}.pkl'.format(score_name)))

    # save as csv
    df_analysis.to_csv(os.path.join(post_eval_path, 'analysis_{}.csv'.format(score_name)))

    print('done')


if __name__ == '__main__':
    state = 0  # [0 - 'eval_before_training', 1 - 'eval_after_training']
    run_post_eval = True  # set True for regular running - post_eval + analysis

    # path to evaluate pickle
    before_folder = training_args.path_args.pretraining_folder
    after_folder = training_args.path_args.after_training_folder
    results_folder = training_args.path_args.results_path

    pickle_list = []
    pickle_names = []
    new_pickle_path_list = []
    if state == 1:
        path_to_predictions = private_args.path.path_to_predictions_after_training
        # iterate over folders in path_to_predictions
        for folder in os.listdir(path_to_predictions):
            pickle_list.append(os.path.join(path_to_predictions, folder, 'inference_results.pkl'))
            pickle_names.append(folder)
    else: # state 0 - before
        path_to_predictions = private_args.path.path_to_predictions_before_training
        pickle_list.append(os.path.join(path_to_predictions, 'inference_results.pkl'))
        pickle_names.append('before_training')

    if state == 1 and run_post_eval:
        if run_post_eval:
            new_pickles_path = '/home/tok/TRBLLmaker/post_eval/'
            compare_params = ['model', 'prompt_type', 'decode_method']
            score_name = 'total_score'
            # iterate over pickles in new_pickles_path that starts with 'example_to_analysis'
            name_iter = 0
            for new_pickle_path in os.listdir(new_pickles_path):
                if new_pickle_path.startswith('example_to_analysis') and new_pickle_path.endswith('.pkl'):
                    df = pd.read_pickle(new_pickle_path)
                    main_path = private_args.path.main_path
                    curr_name = pickle_names[name_iter]
                    name_iter += 1
                    post_eval_path = os.path.join(main_path, 'post_eval', 'analysis_results', curr_name)
                    # if path does not exist, create it
                    if not os.path.exists(post_eval_path):
                        os.makedirs(post_eval_path)
                    analysis(df, compare_params, score_name, new_pickle_path, post_eval_path)
    else:
        for pickle_path, curr_name in zip(pickle_list, pickle_names):
            main_path = private_args.path.main_path
            post_eval_path = os.path.join(main_path, 'post_eval', 'analysis_results', curr_name)
            # if path does not exist, create it
            if not os.path.exists(post_eval_path):
                os.makedirs(post_eval_path)

            # to run analysis only - please fill new_pickle path!!!!
            if run_post_eval:
                df, new_pickle_path = post_eval(pickle_path)
            else:
                new_pickle_path = '/home/mor.ventura/trbll_maker/post_eval/post_eval/example_to_analysis_inference_results.pkl'
                df = pd.read_pickle(new_pickle_path)

            # --- run analysis -----
            # those params relvant only if run_all = 0 !!!
            compare_params = ['model', 'prompt_type', 'decode_method']
            score_name = 'total_score'
            analysis(df, compare_params, score_name, new_pickle_path, post_eval_path)

# #---------------------------------------------------------------
# # create a doc file to write the generated prompts
# doc = docx.Document()
# doc.add_heading('Predicted annotations compare {}', 0)
#
# # compare the same prompt and decode method with different models
# # print the input prompt and the predicted text for each model
# input_prompt = df.loc[0, 'input_prompt']
# df_input_prompt = df[df['input_prompt'] == input_prompt]
# for index, row in df_input_prompt.iterrows():
#     para = doc.add_paragraph("model: {}\n".format(row['model']))
#     para.add_run("decode method: {}\n".format(row['decode_method']))
#     para.add_run("predicted text: {}\n".format(row['predicted_text'])).font.bold = True
#
#
#
# input_prompt = df.loc[5, 'input_prompt']
# df_input_prompt = df[df['input_prompt'] == input_prompt]
# for index, row in df_input_prompt.iterrows():
#     para = doc.add_paragraph("model: {}\n".format(row['model']))
#     para.add_run("decode method: {}\n".format(row['decode_method']))
#     para.add_run("predicted text: {}\n".format(row['predicted_text'])).font.bold = True
#
# doc.save('/home/tok/TRBLLmaker/results/{}/{}.docx'.format(folder, pickle_name.split('.')[0]))
#
# # compare the same prompt and model with different decode methods
#
# # compare the same model and decode method with different prompts

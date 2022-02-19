from datasets import load_dataset, load_metric
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import docx
from docx.enum.text import WD_COLOR_INDEX
from box import Box
import yaml
import os
import datetime
import re
import string
os.environ['KMP_DUPLICATE_LIB_OK'] = 'True'


def explore_data():
    with open('config.yaml') as f:
        training_args = Box(yaml.load(f, Loader=yaml.FullLoader))

    # Load data
    samples = pd.read_json('jsons/samples_final.json')
    songs = pd.read_json('jsons/songs_final.json')

    # Calculate statistics
    lyrics_length = []
    for songTuple in songs.itertuples():
        lyrics = songTuple.lyrics
        lyrics_length.append(len(lyrics))

    text_length = []
    annotation_length = []
    for sampleTuple in samples.itertuples():
        text = sampleTuple.text
        annotation = sampleTuple.annotation
        text_length.append(len(text))
        annotation_length.append(len(annotation))

    # Plot statistics
    # Genre distribution
    songs.genre.value_counts().sort_values().plot(kind='barh')
    plt.title('Distribution of songs genres')
    plt.ylabel('Genre')
    plt.xlabel('Number of songs')
    plt.show()
    # Artist distribution
    artist_hist = songs.artist.value_counts().sort_values(ascending=False)
    artist_hist.head(20).plot(kind='barh')
    plt.title('Distribution of songs artists')
    plt.ylabel('Artist')
    plt.xlabel('Number of songs')
    plt.show()
    # Song length histogram
    plt.hist(lyrics_length, bins=100, range=(0, 10000), density=True,
             color='blue', edgecolor='black', linewidth=1.2)
    plt.title('Distribution of songs length')
    plt.xlabel('Length')
    plt.ylabel('Number of songs')
    plt.show()
    # Boxplot songs length
    plt.boxplot(lyrics_length)
    plt.title('Boxplot of songs length')
    plt.xlabel('Length')
    plt.ylabel('Number of songs')
    plt.show()
    # Text length distribution (text of the annotation)
    plt.hist(text_length, bins=100, range=(0, 10000), density=True,
             color='blue', edgecolor='black', linewidth=1.2)
    plt.title('Distribution of examples length')
    plt.xlabel('Length')
    plt.ylabel('Number of examples')
    plt.show()
    # Plot Boxplot of songs text length
    plt.boxplot(text_length)
    plt.title('Boxplot of examples length')
    plt.xlabel('Length')
    plt.ylabel('Number of examples')
    plt.show()
    # Annotation length distribution
    plt.hist(annotation_length, bins=100, range=(0, 10000), density=True,
             color='blue', edgecolor='black', linewidth=1.2)
    plt.title('Distribution of annotations length')
    plt.xlabel('Length')
    plt.ylabel('Number of annotations')
    plt.show()
    # Boxplot of annotation length
    plt.boxplot(annotation_length)
    plt.title('Boxplot of annotations length')
    plt.xlabel('Length')
    plt.ylabel('Number of annotations')
    plt.show()

    # print examples
    filename = "examples.docx"
    number_of_samples = 40
    random_samples = np.random.choice(samples.index, number_of_samples, replace=False)
    doc = docx.Document()
    doc.add_heading('Annotations', 0)
    for sample_id in random_samples:
        sample_tuple = samples.loc[sample_id]
        text = sample_tuple.text
        annotation = sample_tuple.annotation
        song_id = sample_tuple.song_id
        para = doc.add_paragraph('Example: ' + str(sample_id) + '\n')
        para.add_run('Text: ' + text + '\n').bold = True
        para.add_run('Annotation: ' + annotation + '\n').font.highlight_color = docx.enum.text.WD_COLOR_INDEX.BRIGHT_GREEN
        song = songs[songs.song_id == song_id].iloc[0]
        para.add_run('Artist: ' + song.artist + '\n').bold = True
        para.add_run('Song: ' + song.lyrics + '\n')
    # Now save the document
    doc.save(filename + ' ' + datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S") + '.docx')


def clean_data():
    with open('config.yaml') as f:
        training_args = Box(yaml.load(f, Loader=yaml.FullLoader))

    # Load data
    samples = pd.read_json('jsons/samples_final.json')
    songs = pd.read_json('jsons/songs_final.json')

    bad_song_ids = []
    char_set = string.ascii_letters + string.digits + string.punctuation + '•' + '–' + '…' + '”' + '“' + '‘' + '’' +\
               '—' + ' ' + '\n' + '\r' + '\t'

    for index, sample in enumerate(samples.itertuples()):
        text = sample.text
        annotation = sample.annotation
        song_id = sample.song_id
        is_bad = False
        for text_char in text:  # check if text contains only allowed characters
            if text_char not in char_set:
                bad_song_ids.append(song_id)
                samples.drop(index, inplace=True)
                is_bad = True
                break
        if is_bad:
            continue
        for annotation_char in annotation:  # check if annotation contains only allowed characters
            if annotation_char not in char_set:
                bad_song_ids.append(song_id)
                samples.drop(index, inplace=True)
                break

    # remove links from annotations
    samples['annotation'] = samples['annotation'].str.replace('http\S+', '')

    samples.to_json('jsons/samples_cleaned.json')
    # samples.to_csv('csvs/samples_cleaned.csv', index=False)

    bad_song_ids = list(set(bad_song_ids))
    songs = songs[~songs.song_id.isin(bad_song_ids)]
    songs.to_json('jsons/songs_cleaned.json')
    # songs.to_csv('csvs/songs_cleaned.csv', index=False)


if __name__ == '__main__':
    clean_data()
    # explore_data()
